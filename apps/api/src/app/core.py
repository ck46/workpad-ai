from __future__ import annotations

import json
import re
from html import escape as html_escape
from io import BytesIO
from dataclasses import dataclass
from datetime import UTC, datetime
from functools import lru_cache
from pathlib import Path
from typing import Any
from uuid import uuid4

from docx import Document
from docx.shared import Pt
from markdown import markdown
from pydantic_settings import BaseSettings, SettingsConfigDict
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer
from sqlalchemy import JSON, DateTime, ForeignKey, Integer, LargeBinary, String, Text, UniqueConstraint, create_engine, func, select
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, relationship, sessionmaker

from .schemas import ArtifactRead, ArtifactUpdateRequest, CanvasToolCall, CitationRead, ContentType, ConversationDetail, ConversationSummary, MessageRead


def utcnow() -> datetime:
    return datetime.now(UTC)


class Settings(BaseSettings):
    app_name: str = "Workpad AI"
    api_prefix: str = "/api"
    openai_api_key: str = ""
    anthropic_api_key: str = ""
    default_model: str = "gpt-5.4"
    openai_reasoning_effort: str = "medium"
    github_default_token: str = ""
    app_database_url: str = "sqlite:///./data/workpad.db"
    cors_origins: str = "http://localhost:3000,http://127.0.0.1:3000,http://host.docker.internal:3000"

    model_config = SettingsConfigDict(env_file=".env", env_file_encoding="utf-8", case_sensitive=False, extra="ignore")

    @property
    def cors_origins_list(self) -> list[str]:
        if not self.cors_origins:
            return ["http://localhost:3000", "http://127.0.0.1:3000", "http://host.docker.internal:3000"]
        return [item.strip() for item in self.cors_origins.split(",") if item.strip()]


@lru_cache
def get_settings() -> Settings:
    return Settings()


class Base(DeclarativeBase):
    pass


class Conversation(Base):
    __tablename__ = "conversations"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    title: Mapped[str] = mapped_column(String(240), default="New workpad")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow, onupdate=utcnow)
    archived_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True, default=None)

    messages: Mapped[list["Message"]] = relationship(back_populates="conversation", cascade="all, delete-orphan")
    artifacts: Mapped[list["Artifact"]] = relationship(back_populates="conversation", cascade="all, delete-orphan")


class Message(Base):
    __tablename__ = "messages"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    conversation_id: Mapped[str] = mapped_column(ForeignKey("conversations.id", ondelete="CASCADE"), index=True)
    role: Mapped[str] = mapped_column(String(32))
    content: Mapped[str] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)

    conversation: Mapped[Conversation] = relationship(back_populates="messages")


class Artifact(Base):
    __tablename__ = "artifacts"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    conversation_id: Mapped[str] = mapped_column(ForeignKey("conversations.id", ondelete="CASCADE"), index=True)
    title: Mapped[str] = mapped_column(String(240))
    content: Mapped[str] = mapped_column(Text, default="")
    content_type: Mapped[str] = mapped_column(String(32), default=ContentType.MARKDOWN.value)
    spec_type: Mapped[str | None] = mapped_column(String(32), nullable=True, default=None)
    version: Mapped[int] = mapped_column(Integer, default=1)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow, onupdate=utcnow)

    conversation: Mapped[Conversation] = relationship(back_populates="artifacts")
    versions: Mapped[list["ArtifactVersion"]] = relationship(back_populates="artifact", cascade="all, delete-orphan")


class SpecSource(Base):
    __tablename__ = "spec_sources"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    artifact_id: Mapped[str] = mapped_column(ForeignKey("artifacts.id", ondelete="CASCADE"), index=True)
    kind: Mapped[str] = mapped_column(String(32))
    payload: Mapped[dict] = mapped_column(JSON)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)

    artifact: Mapped[Artifact] = relationship()


class Citation(Base):
    __tablename__ = "citations"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    artifact_id: Mapped[str] = mapped_column(ForeignKey("artifacts.id", ondelete="CASCADE"), index=True)
    anchor: Mapped[str] = mapped_column(String(64), index=True)
    kind: Mapped[str] = mapped_column(String(32))
    target: Mapped[dict] = mapped_column(JSON)
    resolved_state: Mapped[str] = mapped_column(String(16), default="unknown")
    last_checked_at: Mapped[datetime | None] = mapped_column(DateTime(timezone=True), nullable=True, default=None)
    last_observed: Mapped[dict | None] = mapped_column(JSON, nullable=True, default=None)
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)

    artifact: Mapped[Artifact] = relationship()


class RepoCache(Base):
    __tablename__ = "repo_cache"
    __table_args__ = (
        UniqueConstraint("repo", "ref", "path", name="uq_repo_cache_repo_ref_path"),
    )

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    repo: Mapped[str] = mapped_column(String(240), index=True)
    ref: Mapped[str] = mapped_column(String(64))
    path: Mapped[str] = mapped_column(String(1024))
    content: Mapped[bytes] = mapped_column(LargeBinary)
    content_hash: Mapped[str] = mapped_column(String(128))
    etag: Mapped[str | None] = mapped_column(String(256), nullable=True, default=None)
    fetched_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)


class ArtifactVersion(Base):
    __tablename__ = "artifact_versions"

    id: Mapped[str] = mapped_column(String(36), primary_key=True, default=lambda: str(uuid4()))
    artifact_id: Mapped[str] = mapped_column(ForeignKey("artifacts.id", ondelete="CASCADE"), index=True)
    version: Mapped[int] = mapped_column(Integer)
    title: Mapped[str] = mapped_column(String(240))
    content: Mapped[str] = mapped_column(Text)
    content_type: Mapped[str] = mapped_column(String(32))
    summary: Mapped[str] = mapped_column(String(500), default="")
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=utcnow)

    artifact: Mapped[Artifact] = relationship(back_populates="versions")


@lru_cache
def get_engine():
    settings = get_settings()
    database_url = settings.app_database_url
    if database_url.startswith("sqlite:///"):
        db_path = Path(database_url.replace("sqlite:///", "", 1))
        db_path.parent.mkdir(parents=True, exist_ok=True)
        return create_engine(database_url, connect_args={"check_same_thread": False})
    return create_engine(database_url)


@lru_cache
def get_session_factory():
    return sessionmaker(bind=get_engine(), autoflush=False, autocommit=False, expire_on_commit=False)


def init_db() -> None:
    engine = get_engine()
    Base.metadata.create_all(bind=engine)
    _ensure_conversation_schema(engine)
    _ensure_artifact_schema(engine)


def _ensure_conversation_schema(engine) -> None:
    if engine.dialect.name != "sqlite":
        return
    with engine.begin() as connection:
        columns = {
            row[1]
            for row in connection.exec_driver_sql("PRAGMA table_info('conversations')").fetchall()
        }
        if "archived_at" not in columns:
            connection.exec_driver_sql("ALTER TABLE conversations ADD COLUMN archived_at DATETIME")


def _ensure_artifact_schema(engine) -> None:
    if engine.dialect.name != "sqlite":
        return
    with engine.begin() as connection:
        columns = {
            row[1]
            for row in connection.exec_driver_sql("PRAGMA table_info('artifacts')").fetchall()
        }
        if "spec_type" not in columns:
            connection.exec_driver_sql("ALTER TABLE artifacts ADD COLUMN spec_type VARCHAR(32)")


def serialize_conversation(conversation: Conversation, session: Session) -> ConversationSummary:
    last_message = session.scalar(
        select(Message).where(Message.conversation_id == conversation.id).order_by(Message.created_at.desc()).limit(1)
    )
    artifact_count = session.scalar(
        select(func.count()).select_from(Artifact).where(Artifact.conversation_id == conversation.id)
    ) or 0
    preview = None
    if last_message:
        preview = last_message.content[:120]
    return ConversationSummary(
        id=conversation.id,
        title=conversation.title,
        created_at=conversation.created_at,
        updated_at=conversation.updated_at,
        last_message_preview=preview,
        artifact_count=int(artifact_count),
        archived_at=conversation.archived_at,
    )


def serialize_message(message: Message) -> MessageRead:
    return MessageRead(
        id=message.id,
        role=message.role,
        content=message.content,
        created_at=message.created_at,
    )


def serialize_citation(citation: Citation) -> CitationRead:
    return CitationRead(
        id=citation.id,
        artifact_id=citation.artifact_id,
        anchor=citation.anchor,
        kind=citation.kind,
        target=citation.target or {},
        resolved_state=citation.resolved_state,
        last_checked_at=citation.last_checked_at,
        last_observed=citation.last_observed,
    )


def serialize_artifact(artifact: Artifact, session: Session | None = None) -> ArtifactRead:
    citations: list[CitationRead] = []
    if session is not None:
        rows = session.scalars(
            select(Citation).where(Citation.artifact_id == artifact.id).order_by(Citation.created_at.asc())
        ).all()
        citations = [serialize_citation(row) for row in rows]
    return ArtifactRead(
        id=artifact.id,
        conversation_id=artifact.conversation_id,
        title=artifact.title,
        content=artifact.content,
        content_type=artifact.content_type,
        version=artifact.version,
        updated_at=artifact.updated_at,
        citations=citations,
    )


def list_conversations(session: Session, *, include_archived: bool = False) -> list[ConversationSummary]:
    query = select(Conversation).order_by(Conversation.updated_at.desc())
    if not include_archived:
        query = query.where(Conversation.archived_at.is_(None))
    conversations = session.scalars(query).all()
    return [serialize_conversation(item, session) for item in conversations]


def archive_conversation(session: Session, conversation_id: str) -> Conversation:
    conversation = get_conversation_or_404(session, conversation_id)
    if conversation.archived_at is None:
        conversation.archived_at = utcnow()
        conversation.updated_at = utcnow()
        session.commit()
        session.refresh(conversation)
    return conversation


def unarchive_conversation(session: Session, conversation_id: str) -> Conversation:
    conversation = get_conversation_or_404(session, conversation_id)
    if conversation.archived_at is not None:
        conversation.archived_at = None
        conversation.updated_at = utcnow()
        session.commit()
        session.refresh(conversation)
    return conversation


def delete_conversation(session: Session, conversation_id: str) -> None:
    conversation = get_conversation_or_404(session, conversation_id)
    session.delete(conversation)
    session.commit()


def create_conversation(session: Session, seed_title: str | None = None) -> Conversation:
    title = (seed_title or "New workpad").strip() or "New workpad"
    conversation = Conversation(title=title[:240])
    session.add(conversation)
    session.commit()
    session.refresh(conversation)
    return conversation


def get_conversation_or_404(session: Session, conversation_id: str) -> Conversation:
    conversation = session.get(Conversation, conversation_id)
    if conversation is None:
        raise ValueError("Conversation not found")
    return conversation


def get_conversation_detail(session: Session, conversation_id: str) -> ConversationDetail:
    conversation = get_conversation_or_404(session, conversation_id)
    messages = session.scalars(select(Message).where(Message.conversation_id == conversation_id).order_by(Message.created_at.asc())).all()
    artifacts = session.scalars(select(Artifact).where(Artifact.conversation_id == conversation_id).order_by(Artifact.updated_at.desc())).all()
    active_artifact_id = artifacts[0].id if artifacts else None
    return ConversationDetail(
        conversation=serialize_conversation(conversation, session),
        messages=[serialize_message(message) for message in messages],
        artifacts=[serialize_artifact(artifact) for artifact in artifacts],
        active_artifact_id=active_artifact_id,
    )


def add_message(session: Session, conversation: Conversation, role: str, content: str) -> Message:
    message = Message(conversation_id=conversation.id, role=role, content=content)
    session.add(message)
    if conversation.title == "New workpad" and role == "user":
        conversation.title = content.strip().splitlines()[0][:80]
    conversation.updated_at = utcnow()
    session.commit()
    session.refresh(message)
    session.refresh(conversation)
    return message


def get_last_message_by_role(session: Session, conversation: Conversation, role: str) -> Message | None:
    return session.scalar(
        select(Message)
        .where(Message.conversation_id == conversation.id, Message.role == role)
        .order_by(Message.created_at.desc(), Message.id.desc())
        .limit(1)
    )


def delete_messages_after(session: Session, conversation: Conversation, reference: Message) -> None:
    trailing = session.scalars(
        select(Message)
        .where(
            Message.conversation_id == conversation.id,
            Message.created_at > reference.created_at,
        )
        .order_by(Message.created_at.asc())
    ).all()
    for message in trailing:
        if message.id == reference.id:
            continue
        session.delete(message)
    if trailing:
        conversation.updated_at = utcnow()
        session.commit()
        session.refresh(conversation)


def prepare_regenerate(session: Session, conversation: Conversation) -> Message:
    last_user = get_last_message_by_role(session, conversation, "user")
    if last_user is None:
        raise ValueError("No user message to regenerate from.")
    delete_messages_after(session, conversation, last_user)
    return last_user


def apply_edit_to_last_user(session: Session, conversation: Conversation, new_content: str) -> Message:
    last_user = get_last_message_by_role(session, conversation, "user")
    if last_user is None:
        raise ValueError("No user message to edit.")
    cleaned = new_content.strip()
    if not cleaned:
        raise ValueError("Edited message cannot be empty.")
    delete_messages_after(session, conversation, last_user)
    last_user.content = cleaned
    conversation.updated_at = utcnow()
    session.commit()
    session.refresh(last_user)
    session.refresh(conversation)
    return last_user


@dataclass
class ArtifactMutationResult:
    artifact: Artifact
    summary: str
    action: str


def _record_artifact_version(session: Session, artifact: Artifact, summary: str) -> None:
    session.add(
        ArtifactVersion(
            artifact_id=artifact.id,
            version=artifact.version,
            title=artifact.title,
            content=artifact.content,
            content_type=artifact.content_type,
            summary=summary[:500],
        )
    )


def _apply_patches(content: str, tool_call: CanvasToolCall) -> str:
    next_content = content
    for patch in tool_call.patches or []:
        if patch.search not in next_content:
            if patch.allow_missing:
                continue
            raise ValueError(f"Could not find patch target: {patch.search[:80]}")
        count = -1 if patch.replace_all else 1
        next_content = next_content.replace(patch.search, patch.replace, count)
    return next_content


def apply_canvas_tool(
    session: Session,
    conversation: Conversation,
    tool_call: CanvasToolCall,
    current_artifact_id: str | None = None,
) -> ArtifactMutationResult:
    artifact: Artifact | None = None
    if current_artifact_id:
        artifact = session.get(Artifact, current_artifact_id)

    if artifact is None and tool_call.action != "create":
        artifact = session.scalar(
            select(Artifact).where(Artifact.conversation_id == conversation.id).order_by(Artifact.updated_at.desc()).limit(1)
        )

    if artifact is None:
        artifact = Artifact(
            conversation_id=conversation.id,
            title=tool_call.title,
            content="",
            content_type=tool_call.content_type.value,
            version=0,
        )
        session.add(artifact)
        session.flush()

    current_content = artifact.content or ""
    if tool_call.action == "patch":
        next_content = _apply_patches(current_content, tool_call)
    else:
        next_content = tool_call.content or ""

    artifact.title = tool_call.title[:240]
    artifact.content = next_content
    artifact.content_type = tool_call.content_type.value
    artifact.version += 1
    artifact.updated_at = utcnow()
    conversation.updated_at = utcnow()

    _record_artifact_version(session, artifact, tool_call.summary)
    session.commit()
    session.refresh(artifact)
    session.refresh(conversation)

    return ArtifactMutationResult(artifact=artifact, summary=tool_call.summary, action=tool_call.action)


def update_artifact_manually(session: Session, artifact_id: str, payload: ArtifactUpdateRequest) -> ArtifactRead:
    artifact = session.get(Artifact, artifact_id)
    if artifact is None:
        raise ValueError("Artifact not found")
    if payload.expected_version and artifact.version != payload.expected_version:
        raise ValueError("Artifact version mismatch")

    artifact.title = payload.title[:240]
    artifact.content = payload.content
    artifact.content_type = payload.content_type.value
    artifact.version += 1
    artifact.updated_at = utcnow()
    artifact.conversation.updated_at = utcnow()
    _record_artifact_version(session, artifact, "Manual edit from the workpad editor.")
    session.commit()
    session.refresh(artifact)
    return serialize_artifact(artifact)


def get_artifact_or_404(session: Session, artifact_id: str) -> Artifact:
    artifact = session.get(Artifact, artifact_id)
    if artifact is None:
        raise ValueError("Artifact not found")
    return artifact


def _safe_filename(name: str) -> str:
    base = re.sub(r"[^\w\s.-]", "", name).strip()
    base = re.sub(r"\s+", "_", base)
    return base or "workpad_export"


def _iter_markdown_blocks(content: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    buffer: list[str] = []
    in_code_block = False

    def flush_paragraph() -> None:
        nonlocal buffer
        if buffer:
            blocks.append(("paragraph", "\n".join(buffer).strip()))
            buffer = []

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                blocks.append(("code", "\n".join(buffer)))
                buffer = []
                in_code_block = False
            else:
                flush_paragraph()
                in_code_block = True
                buffer = []
            continue

        if in_code_block:
            buffer.append(line)
            continue

        if not stripped:
            flush_paragraph()
            blocks.append(("spacer", ""))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            blocks.append((f"heading_{min(level, 4)}", heading_match.group(2).strip()))
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            blocks.append(("bullet", bullet_match.group(1).strip()))
            continue

        numbered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if numbered_match:
            flush_paragraph()
            blocks.append(("numbered", numbered_match.group(1).strip()))
            continue

        buffer.append(stripped)

    if in_code_block and buffer:
        blocks.append(("code", "\n".join(buffer)))
    elif buffer:
        blocks.append(("paragraph", "\n".join(buffer).strip()))

    return [block for block in blocks if block[0] != "spacer" or block[1]]


def _build_docx_bytes(artifact: Artifact) -> bytes:
    document = Document()
    document.add_heading(artifact.title, level=0)

    for kind, text in _iter_markdown_blocks(artifact.content):
        if kind.startswith("heading_"):
            level = int(kind.rsplit("_", 1)[1])
            document.add_heading(text, level=level)
            continue
        if kind == "bullet":
            document.add_paragraph(text, style="List Bullet")
            continue
        if kind == "numbered":
            document.add_paragraph(text, style="List Number")
            continue
        if kind == "code":
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text or " ")
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue
        if kind == "paragraph":
            document.add_paragraph(text)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_pdf_bytes(artifact: Artifact) -> bytes:
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title=artifact.title,
    )
    styles = getSampleStyleSheet()
    normal = styles["BodyText"]
    heading_styles = {
        "heading_1": styles["Heading1"],
        "heading_2": styles["Heading2"],
        "heading_3": styles["Heading3"],
        "heading_4": ParagraphStyle("Heading4Custom", parent=styles["Heading3"], fontSize=12, leading=15),
    }
    code_style = ParagraphStyle(
        "CodeBlock",
        parent=normal,
        fontName="Courier",
        fontSize=9,
        leading=12,
        leftIndent=12,
        backColor="#F4F4F5",
    )

    story = [Paragraph(artifact.title, styles["Title"]), Spacer(1, 0.2 * inch)]

    for kind, text in _iter_markdown_blocks(artifact.content):
        if kind in heading_styles:
            story.append(Paragraph(text, heading_styles[kind]))
        elif kind == "bullet":
            story.append(Paragraph(f"• {text}", normal))
        elif kind == "numbered":
            story.append(Paragraph(text, normal))
        elif kind == "code":
            story.append(Preformatted(text or " ", code_style))
        elif kind == "paragraph":
            story.append(Paragraph(text.replace("\n", "<br/>"), normal))
        story.append(Spacer(1, 0.12 * inch))

    document.build(story)
    return output.getvalue()


def export_artifact(session: Session, artifact_id: str, export_format: str) -> tuple[bytes | str, str, str]:
    artifact = get_artifact_or_404(session, artifact_id)
    safe_name = _safe_filename(artifact.title)
    if export_format == "html":
        if artifact.content_type == ContentType.MARKDOWN.value:
            body = markdown(artifact.content, extensions=["fenced_code", "tables"])
        else:
            body = f"<pre>{html_escape(artifact.content)}</pre>"
        return body, "text/html; charset=utf-8", f"{safe_name}.html"
    if export_format == "text":
        return artifact.content, "text/plain; charset=utf-8", f"{safe_name}.txt"
    if export_format == "docx":
        return _build_docx_bytes(artifact), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{safe_name}.docx"
    if export_format == "pdf":
        return _build_pdf_bytes(artifact), "application/pdf", f"{safe_name}.pdf"
    return artifact.content, "text/markdown; charset=utf-8", f"{safe_name}.md"


def current_artifact_id_from_payload(payload: Any) -> str | None:
    if not payload:
        return None
    artifact_id = getattr(payload, "id", None)
    if artifact_id:
        return artifact_id
    if isinstance(payload, dict):
        return payload.get("id")
    return None


def build_response_input(session: Session, conversation: Conversation) -> list[dict[str, str]]:
    history = session.scalars(select(Message).where(Message.conversation_id == conversation.id).order_by(Message.created_at.asc())).all()
    return [{"role": message.role, "content": message.content} for message in history]


def json_dumps(data: Any) -> str:
    return json.dumps(data, ensure_ascii=True)
